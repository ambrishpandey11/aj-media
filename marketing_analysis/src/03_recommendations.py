"""
03_recommendations.py — Task 3: Strategic Recommendations
Outputs:
  - outputs/reports/executive_summary.docx
  - outputs/figures/budget_reallocation.png
  - outputs/figures/ab_test_plan.png
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
from utils import *

# ─────────────────────────────────────────────────────────────────────────────
# BUDGET REALLOCATION FRAMEWORK (30% cut from bottom)
# ─────────────────────────────────────────────────────────────────────────────
def build_reallocation_plan(df: pd.DataFrame) -> dict:
    """Identify bottom 30% spend segments and reallocation targets."""
    # Rank every (platform × region × creative) combo by ROAS
    combo = df.groupby(["Platform", "Region", "Creative_Type"]).agg(
        ROAS=("ROAS","mean"), Spend=("Spend","sum"),
        Revenue=("Revenue","sum"), Purchases=("Purchases","sum")
    ).reset_index().sort_values("ROAS")

    total_spend = combo["Spend"].sum()
    target_cut  = total_spend * 0.30

    # Walk from worst ROAS upward, accumulate cuts until 30% reached
    combo["cum_spend"] = combo["Spend"].cumsum()
    cut_mask = combo["cum_spend"] <= target_cut
    cut_df   = combo[cut_mask].copy()
    keep_df  = combo[~cut_mask].copy()
    cut_df["recommendation"] = "CUT / Pause"
    keep_df["recommendation"] = "Maintain / Scale"

    # Flag highest ROAS for reallocation
    top10 = keep_df.nlargest(10, "ROAS")
    top10["recommendation"] = "Scale Up"
    keep_df.loc[top10.index, "recommendation"] = "Scale Up"

    actual_cut = cut_df["Spend"].sum()
    print(f"\n  Budget Reallocation Plan:")
    print(f"    Total spend:        ${total_spend:>12,.0f}")
    print(f"    Target cut (30%):   ${target_cut:>12,.0f}")
    print(f"    Actual cut:         ${actual_cut:>12,.0f} ({actual_cut/total_spend*100:.1f}%)")
    print(f"    Segments to cut:    {len(cut_df)}")
    print(f"    Segments to scale:  {len(top10)}")

    return {"cut": cut_df, "keep": keep_df, "top10": top10,
            "total_spend": total_spend, "cut_amount": actual_cut}

def plot_budget_reallocation(df: pd.DataFrame, plan: dict):
    set_style()
    fig, axes = plt.subplots(1, 3, figsize=(18, 7))

    # Current vs proposed spend by platform
    ax = axes[0]
    plat = df.groupby("Platform")["Spend"].sum()
    cut_by_plat = plan["cut"].groupby("Platform")["Spend"].sum().reindex(plat.index, fill_value=0)
    remaining = plat - cut_by_plat

    x = np.arange(len(plat))
    w = 0.35
    ax.bar(x - w/2, plat / 1e6, w, label="Current",  color=COLORS["danger"],  alpha=0.8)
    ax.bar(x + w/2, remaining / 1e6, w, label="Proposed", color=COLORS["ok"], alpha=0.8)
    ax.set_xticks(x); ax.set_xticklabels(plat.index)
    ax.set_title("Current vs Proposed Spend\nby Platform ($M)", fontweight="bold")
    ax.set_ylabel("Spend ($M)"); ax.legend(fontsize=8)

    # ROAS waterfall: before vs expected after
    ax = axes[1]
    current_roas = df["ROAS"].mean()
    # Expected improvement: removing low-ROAS spend raises average
    weighted_new = (plan["keep"]["ROAS"] * plan["keep"]["Spend"]).sum() / plan["keep"]["Spend"].sum()
    categories = ["Current\nROAS", "After Cut\n(estimate)", "Target\nROAS", "Industry\nBenchmark"]
    values     = [current_roas, weighted_new, 1.2, 1.4]
    colors_r   = [roas_color(v) for v in values]
    bars = ax.bar(categories, values, color=colors_r, edgecolor="white", width=0.5)
    ax.set_title("ROAS: Current → Expected\nAfter Reallocation", fontweight="bold")
    ax.set_ylabel("ROAS")
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.02,
                f"{val:.2f}", ha="center", fontsize=10, fontweight="bold")

    # Heatmap: ROAS by Platform × Region
    ax = axes[2]
    pivot = df.pivot_table(values="ROAS", index="Region", columns="Platform", aggfunc="mean")
    sns.heatmap(pivot, ax=ax, annot=True, fmt=".2f",
                cmap="RdYlGn", center=1.2, linewidths=0.5,
                cbar_kws={"label": "ROAS"})
    ax.set_title("ROAS Heatmap: Cut/Scale Guide\n(Red = cut, Green = scale)", fontweight="bold")

    fig.suptitle("Budget Reallocation Framework — 30% Cut Strategy", fontweight="bold", fontsize=13)
    save_fig("budget_reallocation")

# ─────────────────────────────────────────────────────────────────────────────
# GENERATE WORD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
def generate_word_report(df: pd.DataFrame, plan: dict):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  [warn] python-docx not installed — skipping Word doc. Run: pip install python-docx")
        return

    print("\n  Generating executive_summary.docx...")
    doc = Document()

    # ── Styles ────────────────────────────────────────────────────────────────
    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        if level == 1:
            p.runs[0].font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
        return p

    def para(text, bold=False, italic=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        if color: run.font.color.rgb = color
        return p

    def bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        return p

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Shading Accent 1"
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for row_data in rows:
            row = table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val)
        return table

    # ── Key numbers ──────────────────────────────────────────────────────────
    total_spend    = df["Spend"].sum()
    total_rev      = df["Revenue"].sum()
    overall_roas   = total_rev / total_spend if total_spend else 0
    best_platform  = df.groupby("Platform")["ROAS"].mean().idxmax()
    worst_platform = df.groupby("Platform")["ROAS"].mean().idxmin()
    best_region    = df.groupby("Region")["ROAS"].mean().idxmax()
    best_product   = df.groupby("Product_Category")["ROAS"].mean().idxmax()
    best_creative  = df.groupby("Creative_Type")["ROAS"].mean().idxmax()
    cut_amount     = plan["cut_amount"]

    # ── Title page ───────────────────────────────────────────────────────────
    doc.add_heading("DTC Fitness Supplements", 0)
    heading("Marketing Performance: Executive Summary & Strategic Recommendations")
    para(f"Period: January – March 2024  |  Prepared by: Data Analytics Team", italic=True)
    doc.add_paragraph()

    # ── Section 1: Key Findings ───────────────────────────────────────────────
    heading("1. Key Insights from Analysis", 1)

    findings = [
        (f"Critical ROAS Gap",
         f"Overall ROAS of {overall_roas:.2f} is {(1.2-overall_roas)/1.2*100:.0f}% below the 1.2 target and "
         f"{(1.4-overall_roas)/1.4*100:.0f}% below the 1.4 industry benchmark. "
         f"Every $1 spent generates only ${overall_roas:.2f} in revenue."),
        (f"Platform Performance Divergence",
         f"{best_platform} leads with the highest ROAS while {worst_platform} consistently underperforms. "
         f"Spend is not optimally allocated to mirror these returns."),
        (f"Regional Opportunity",
         f"{best_region} shows the strongest regional ROAS. Budget currently distributed roughly evenly "
         f"across regions despite significant performance differences."),
        (f"Creative Type Signal",
         f"{best_creative} creative achieves the highest ROAS. Video completion rate shows a positive "
         f"correlation with CVR — high-completion videos are strong conversion drivers."),
        (f"Competitive Event Impact",
         f"CPM increases during competitive periods while CVR declines, suggesting the brand loses "
         f"bidding efficiency during peak competitor activity. Bidding strategy should be adjusted."),
    ]

    for title, detail in findings:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: "); run.bold = True
        p.add_run(detail)

    # ── Section 2: Immediate Optimisations ────────────────────────────────────
    heading("2. Immediate Optimization Opportunities", 1)

    heading("2a. 30% Spend Reduction Framework", 2)
    para(f"Target cut: ${cut_amount:,.0f} from lowest-performing campaign segments.")
    doc.add_paragraph()

    top_cuts = plan["cut"].head(10)[["Platform","Region","Creative_Type","ROAS","Spend"]]
    top_cuts["ROAS"]  = top_cuts["ROAS"].round(2)
    top_cuts["Spend"] = top_cuts["Spend"].apply(lambda x: f"${x:,.0f}")
    add_table(
        ["Platform","Region","Creative Type","ROAS","Spend"],
        top_cuts.values.tolist()
    )
    doc.add_paragraph()
    para("These segments should be paused or eliminated immediately. Budget freed should be "
         "reallocated to top-performing segments identified below.", italic=True)

    heading("2b. Scale-Up Targets", 2)
    top_scale = plan["top10"][["Platform","Region","Creative_Type","ROAS","Spend"]].head(8)
    top_scale["ROAS"]  = top_scale["ROAS"].round(2)
    top_scale["Spend"] = top_scale["Spend"].apply(lambda x: f"${x:,.0f}")
    add_table(
        ["Platform","Region","Creative Type","ROAS","Current Spend"],
        top_scale.values.tolist()
    )
    doc.add_paragraph()

    heading("2c. Competitive Event Bidding Strategy", 2)
    for item in [
        "Reduce bids by 15–20% during identified competitive event weeks to preserve margin.",
        "Shift budget to Search/Intent-based creatives during competitor surge periods.",
        "Increase retargeting spend during competitive events — own audience converts at higher rate.",
        "Set automated rules: if CPM > 2× 30-day baseline, reduce daily budget by 25%.",
    ]:
        bullet(item)

    # ── Section 3: A/B Test Recommendations ──────────────────────────────────
    heading("3. A/B Testing Recommendations", 1)

    tests = [
        {
            "name": "Test 1: Video Creative Completion Optimisation",
            "hypothesis": "Shorter video ads (≤15s) with strong hooks in the first 3 seconds will increase "
                          "video completion rate and CVR by 15–20% vs. current video length mix.",
            "methodology": "Split audiences 50/50 within same ad sets. Control: existing video creative. "
                           "Treatment: re-edited 15s hook-first versions of top-performing videos.",
            "metrics": "Primary: Video Completion Rate, CVR. Secondary: ROAS, CPC.",
            "impact": "If CVR improves 15%, estimated additional revenue: "
                      f"${df[df['Creative_Type'].str.lower()=='video']['Revenue'].sum() * 0.15:,.0f}/quarter.",
            "timeline": "Week 1–2: Production. Week 3–6: Live test (min 1,000 clicks/arm). Week 7: Analysis.",
        },
        {
            "name": "Test 2: Regional Budget Concentration",
            "hypothesis": f"Concentrating 40% of budget in {best_region} (highest ROAS region) while reducing "
                          "lowest-ROAS region by 40% will improve blended ROAS by 0.1–0.15 points.",
            "methodology": "Implement for 4 weeks. Maintain creative/audience mix identical — change only "
                           "geographic budget weights. Control: current allocation. Treatment: concentration.",
            "metrics": "Primary: Blended ROAS. Secondary: Revenue, total purchases.",
            "impact": "Estimated ROAS improvement of 0.10–0.15 → $150K–$225K additional revenue/month.",
            "timeline": "Week 1: Setup targeting. Week 2–5: Live test. Week 6: Readout.",
        },
        {
            "name": "Test 3: Frequency Cap Optimisation",
            "hypothesis": "Capping ad frequency at 3 exposures/user/week (vs uncapped) will reduce CPM "
                          "fatigue and improve CVR by 10% through audience freshness.",
            "methodology": "Campaign-level test. Control: no frequency cap. Treatment: hard cap at 3/week. "
                           "Same budget, creative, targeting. Run 3 weeks minimum.",
            "metrics": "Primary: CVR, CPC. Secondary: CTR, ROAS, frequency distribution.",
            "impact": "10% CVR improvement on $1.6M/month spend = ~$192K additional monthly revenue.",
            "timeline": "Week 1: Config & QA. Week 2–4: Live. Week 5: Analysis & rollout decision.",
        },
    ]

    for test in tests:
        heading(test["name"], 2)
        for key in ["hypothesis","methodology","metrics","impact","timeline"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{key.capitalize()}: "); run.bold = True
            p.add_run(test[key])

    # ── Section 4: Long-term Strategy ────────────────────────────────────────
    heading("4. Long-term Strategic Recommendations", 1)

    heading("4a. Regional Budget Allocation Strategy", 2)
    para("Proposed quarterly allocation — data-driven reweighting based on ROAS:")
    roas_by_region = df.groupby("Region")["ROAS"].mean()
    region_weight  = (roas_by_region / roas_by_region.sum())
    monthly_budget = 1_600_000
    alloc_table = [(r, f"{roas_by_region[r]:.2f}", f"{region_weight[r]*100:.1f}%",
                    f"${region_weight[r]*monthly_budget:,.0f}")
                   for r in roas_by_region.index]
    add_table(["Region","Avg ROAS","Proposed Share","Monthly Budget"], alloc_table)

    heading("4b. Creative Strategy by Platform", 2)
    creative_by_plat = df.groupby(["Platform","Creative_Type"])["ROAS"].mean().reset_index()
    best_creative_per_plat = creative_by_plat.loc[
        creative_by_plat.groupby("Platform")["ROAS"].idxmax()]
    for _, row in best_creative_per_plat.iterrows():
        bullet(f"{row['Platform']}: Prioritise {row['Creative_Type']} (ROAS {row['ROAS']:.2f})")

    heading("4c. Product & Audience Priorities", 2)
    prod_roas = df.groupby("Product_Category")["ROAS"].mean().sort_values(ascending=False)
    aud_roas  = df.groupby("Target_Audience")["ROAS"].mean().sort_values(ascending=False)
    para("Top product categories by ROAS:")
    for p_name, roas_val in prod_roas.items():
        bullet(f"{p_name}: {roas_val:.2f} ROAS")
    para("Top audience segments by ROAS:")
    for a_name, roas_val in aud_roas.items():
        bullet(f"{a_name}: {roas_val:.2f} ROAS")

    heading("4d. Quarterly Planning Framework", 2)
    for item in [
        "Month 1: Execute 30% spend cut, launch A/B tests 1 & 2.",
        "Month 2: Reallocate freed budget to top-performing segments. Launch A/B test 3.",
        "Month 3: Read A/B results, scale winners. Apply ROAS-weighted regional allocation.",
        "Competitive events: 2 weeks pre-event, shift to retargeting. Post-event: aggressive acquisition.",
        "Monthly review: ROAS by segment report. Pause any segment dropping below 0.9 ROAS.",
    ]:
        bullet(item)

    # ── Section 5: Risk & Measurement ────────────────────────────────────────
    heading("5. Risk Assessment & Measurement Plan", 1)

    heading("Risks", 2)
    risks = [
        ("Revenue drop from cutting spend", "Medium", "Monitor daily revenue for first 2 weeks; rollback trigger if revenue drops >15%"),
        ("A/B test contamination", "Low", "Use campaign-level separation; avoid shared audiences"),
        ("Seasonal ROAS fluctuation", "Medium", "Adjust targets in Q2/Q3 based on historical seasonality"),
        ("Competitor escalation response", "Low–Medium", "Pre-define max CPM thresholds per platform; auto-pause rules"),
    ]
    add_table(["Risk","Likelihood","Mitigation"], risks)

    heading("KPIs to Track Weekly", 2)
    for kpi in ["ROAS (target: ≥1.2 within 60 days)", "CTR and CVR by platform",
                "CPM during competitive events", "Spend by segment vs allocation plan",
                "Revenue vs prior week (WoW change %)", "A/B test statistical significance"]:
        bullet(kpi)

    # Save
    out_path = REPORT_DIR / "executive_summary.docx"
    doc.save(out_path)
    print(f"  [saved] executive_summary.docx")

# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
def run():
    print(f"\n{'='*60}")
    print(f"  TASK 3 — STRATEGIC RECOMMENDATIONS")
    print(f"{'='*60}")

    df = pd.read_csv(REPORT_DIR / "cleaned_data.csv", parse_dates=["Date"])

    plan = build_reallocation_plan(df)
    plot_budget_reallocation(df, plan)
    generate_word_report(df, plan)

    print(f"\n  ✓ Task 3 complete.\n")
    return plan

if __name__ == "__main__":
    run()
